from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path
from uuid import UUID

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session

from app.agents.phase2.output_utils import project_output_dir
from app.db.crud.projects import get_project, list_project_files
from app.db.crud.stages import list_project_stages
from app.db.session import get_db
from app.db.models import (
    AuditEventLog,
    CorrectionEvent,
    Escalation,
    ExtractedFieldValue,
    ParserRun,
    Project,
    ProjectFile,
    Stage,
    ValidationResult,
)
from app.utils.project_paths import get_writable_project_data_root

router = APIRouter()


def _project_output_root(project_id: UUID) -> Path:
    return get_writable_project_data_root() / str(project_id) / "outputs"


def _project_processed_root(project_id: UUID) -> Path:
    return get_writable_project_data_root() / str(project_id) / "Processed"


def _as_download(path: Path) -> dict:
    rel = path.relative_to(path.parents[2]).as_posix() if len(path.parents) > 2 else path.name
    return {
        "name": path.name,
        "relative_path": rel,
        "extension": path.suffix.lower().lstrip(".") or "file",
        "size_bytes": path.stat().st_size,
        "modified_at": datetime.fromtimestamp(path.stat().st_mtime).isoformat(),
    }


def _collect_report_data(db: Session, project_id: UUID) -> dict:
    project = get_project(db, project_id)
    if project is None:
        raise HTTPException(status_code=404, detail="Project not found")

    stages = list_project_stages(db, project_id)
    files = list_project_files(db, project_id)
    output_root = _project_output_root(project_id)
    outputs = []
    if output_root.exists():
        outputs = [
            _as_download(path)
            for path in output_root.rglob("*")
            if path.is_file()
        ]

    return {
        "generated_at": datetime.utcnow().isoformat(),
        "project": {
            "id": str(project.id),
            "proposal_id": project.proposal_id,
            "name": project.name,
            "location": project.location,
            "project_type": project.project_type,
            "status": project.status.value,
        },
        "files": [
            {
                "id": str(file.id),
                "original_filename": file.original_filename,
                "file_type": file.file_type,
                "file_category": file.file_category,
                "source_application": file.source_application,
                "likely_role": file.likely_role,
                "classification_confidence": file.classification_confidence,
                "processing_status": file.processing_status.value,
            }
            for file in files
        ],
        "stages": [
            {
                "stage_code": stage.stage_code,
                "status": stage.status.value,
                "started_at": stage.started_at.isoformat() if stage.started_at else None,
                "completed_at": stage.completed_at.isoformat() if stage.completed_at else None,
                "error_message": stage.error_message,
                "result": json.loads(stage.result_json or "{}"),
            }
            for stage in stages
        ],
        "outputs": outputs,
    }


def _report_lines(data: dict) -> list[str]:
    project = data["project"]
    lines = [
        "Infiniti Steel Detailing Automation",
        "Phase 2 Project Status Report",
        "",
        f"Project: {project['name']}",
        f"Proposal ID: {project['proposal_id']}",
        f"Location: {project.get('location') or '-'}",
        f"Project type: {project.get('project_type') or '-'}",
        f"Generated at: {data['generated_at']}",
        "",
        "Stage Status",
    ]
    for stage in data["stages"]:
        lines.append(f"- {stage['stage_code']}: {stage['status']}")
    lines.extend(["", "File Inventory"])
    for file in data["files"]:
        lines.append(
            f"- {file['original_filename']} | {file['file_type']} | "
            f"{file['likely_role']} | {file['processing_status']} | "
            f"{file['classification_confidence']}%"
        )
    lines.extend(["", "Output Files"])
    if data["outputs"]:
        for output in data["outputs"]:
            lines.append(f"- {output['relative_path']} ({output['size_bytes']} bytes)")
    else:
        lines.append("- No generated output files yet.")
    return lines


def _write_docx_report(path: Path, data: dict) -> None:
    from docx import Document

    doc = Document()
    project = data["project"]
    doc.add_heading("Phase 2 Project Status Report", 0)
    doc.add_paragraph(f"Project: {project['name']}")
    doc.add_paragraph(f"Proposal ID: {project['proposal_id']}")
    doc.add_paragraph(f"Generated at: {data['generated_at']}")

    doc.add_heading("Stage Status", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Stage"
    table.rows[0].cells[1].text = "Status"
    for stage in data["stages"]:
        cells = table.add_row().cells
        cells[0].text = stage["stage_code"]
        cells[1].text = stage["status"]

    doc.add_heading("File Inventory", level=1)
    file_table = doc.add_table(rows=1, cols=5)
    headers = ["File", "Type", "Role", "Status", "Confidence"]
    for index, header in enumerate(headers):
        file_table.rows[0].cells[index].text = header
    for file in data["files"]:
        cells = file_table.add_row().cells
        cells[0].text = file["original_filename"]
        cells[1].text = file["file_type"]
        cells[2].text = file["likely_role"]
        cells[3].text = file["processing_status"]
        cells[4].text = f"{file['classification_confidence']}%"
    doc.save(path)


def _write_pdf_report(path: Path, data: dict) -> None:
    import fitz

    doc = fitz.open()
    page = doc.new_page(width=595, height=842)
    text = "\n".join(_report_lines(data))
    rect = fitz.Rect(42, 42, 553, 800)
    page.insert_textbox(rect, text, fontsize=9, fontname="helv", lineheight=1.25)
    doc.save(path)
    doc.close()


@router.get("/{project_id}/files")
async def list_output_files(project_id: str):
    try:
        pid = UUID(project_id)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail="Invalid project_id") from exc

    output_root = _project_output_root(pid)
    files = []
    if output_root.exists():
        files = [
            {
                **_as_download(path),
                "relative_path": path.relative_to(output_root).as_posix(),
            }
            for path in output_root.rglob("*")
            if path.is_file()
        ]
    return {"project_id": project_id, "files": files}


@router.get("/{project_id}/files/{agent_code}")
async def list_agent_files(project_id: str, agent_code: str):
    """List output files generated by a specific agent (e.g. ab, ga, p2_01)."""
    try:
        pid = UUID(project_id)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail="Invalid project_id") from exc

    # Agent codes like P2-01 might be stored in folders named p2_01
    normalized = agent_code.lower().replace("-", "_")
    agent_dir = project_output_dir(pid, normalized)
    
    files = []
    if agent_dir.exists():
        files = [
            {
                **_as_download(path),
                "relative_path": path.relative_to(agent_dir).as_posix(),
                "agent_code": agent_code
            }
            for path in agent_dir.rglob("*")
            if path.is_file()
        ]
    return {"project_id": project_id, "agent_code": agent_code, "files": files}


@router.get("/{project_id}/download/{file_path:path}")
async def download_output_file(project_id: str, file_path: str):
    try:
        pid = UUID(project_id)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail="Invalid project_id") from exc

    output_root = _project_output_root(pid).resolve()
    target = (output_root / file_path).resolve()
    try:
        target.relative_to(output_root)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail="Invalid output path") from exc
    if not target.exists() or not target.is_file():
        raise HTTPException(status_code=404, detail="Output file not found")
    return FileResponse(target, filename=target.name)


def _write_xlsx_report(path: Path, data: dict) -> None:
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Project Status"
    
    project = data["project"]
    ws["A1"] = "Infiniti Steel Detailing Automation"
    ws["A2"] = "Phase 2 Project Status Report"
    ws["A4"] = "Project Name"
    ws["B4"] = project["name"]
    ws["A5"] = "Proposal ID"
    ws["B5"] = project["proposal_id"]
    ws["A6"] = "Generated At"
    ws["B6"] = data["generated_at"]
    
    ws.append([])
    ws.append(["Stage Status"])
    ws.append(["Stage Code", "Status", "Started At", "Completed At"])
    for s in data["stages"]:
        ws.append([s["stage_code"], s["status"], s["started_at"], s["completed_at"]])
        
    ws.append([])
    ws.append(["File Inventory"])
    ws.append(["Filename", "Type", "Role", "Status", "Confidence (%)"])
    for f in data["files"]:
        ws.append([f["original_filename"], f["file_type"], f["likely_role"], f["processing_status"], f["classification_confidence"]])
        
    wb.save(path)


@router.get("/{project_id}/report/{format_name}")
async def download_phase_report(
    project_id: str,
    format_name: str,
    db: Session = Depends(get_db),
):
    try:
        pid = UUID(project_id)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail="Invalid project_id") from exc

    fmt = format_name.lower()
    if fmt not in {"json", "txt", "docx", "pdf", "xlsx"}:
        raise HTTPException(status_code=400, detail="Supported formats: json, txt, docx, pdf, xlsx")

    data = _collect_report_data(db, pid)
    reports_dir = project_output_dir(pid, "reports")
    report_path = reports_dir / f"phase2_status_report.{fmt}"

    if fmt == "json":
        report_path.write_text(json.dumps(data, indent=2, default=str), encoding="utf-8")
        media_type = "application/json"
    elif fmt == "txt":
        report_path.write_text("\n".join(_report_lines(data)), encoding="utf-8")
        media_type = "text/plain"
    elif fmt == "docx":
        _write_docx_report(report_path, data)
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif fmt == "xlsx":
        _write_xlsx_report(report_path, data)
        media_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    else:
        _write_pdf_report(report_path, data)
        media_type = "application/pdf"

    return FileResponse(report_path, filename=report_path.name, media_type=media_type)



# ---------------------------------------------------------------------------
# Processed/ file helpers
# ---------------------------------------------------------------------------

def _load_processed_json(project_id: UUID, filename: str) -> dict:
    path = _project_processed_root(project_id) / filename
    if not path.exists():
        raise HTTPException(status_code=404, detail=f"Processed file '{filename}' not found")
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Failed to read '{filename}': {exc}") from exc


def _processed_export_dir(project_id: UUID) -> Path:
    path = get_writable_project_data_root() / str(project_id) / "outputs" / "processed_exports"
    path.mkdir(parents=True, exist_ok=True)
    return path


def _query_rows(query) -> list[dict]:
    data = [r.__dict__.copy() for r in query.all()]
    for row in data:
        row.pop("_sa_instance_state", None)
    return data


def _write_json(path: Path, payload: object) -> None:
    path.write_text(json.dumps(payload, indent=2, default=str), encoding="utf-8")


# --- PDF builders ---

def _write_file_inventory_pdf(path: Path, data: dict) -> None:
    import fitz
    doc = fitz.open()
    page = doc.new_page(width=595, height=842)
    y = 42.0

    def _line(text: str, size: float = 9, bold: bool = False, indent: float = 0) -> None:
        nonlocal y
        font = "helv-bold" if bold else "helv"
        page.insert_text(fitz.Point(42 + indent, y), text, fontsize=size, fontname=font)
        y += size * 1.6

    _line("File Inventory Report", 14, bold=True)
    _line(f"Project: {data.get('project_id', '-')}", 9)
    _line(f"Total files: {data.get('total_files', 0)}   Parsed: {data.get('parsed_files', 0)}"
          f"   Failed: {data.get('failed_files', 0)}", 9)
    y += 8

    for f in data.get("files", []):
        if y > 790:
            page = doc.new_page(width=595, height=842)
            y = 42.0
        _line(f"• {f.get('original_filename', '-')}", 9, bold=True)
        _line(f"  Type: {f.get('file_type', '-')}  |  Category: {f.get('file_category', '-')}"
              f"  |  Role: {f.get('likely_role', '-')}", 8, indent=8)
        _line(f"  Source: {f.get('source_application', '-')}"
              f"  |  Status: {f.get('processing_status', '-')}"
              f"  |  Confidence: {f.get('classification_confidence', '-')}", 8, indent=8)
        y += 4

    doc.save(path)
    doc.close()


def _write_governing_source_pdf(path: Path, data: dict) -> None:
    import fitz
    doc = fitz.open()
    page = doc.new_page(width=595, height=842)
    y = 42.0

    def _line(text: str, size: float = 9, bold: bool = False) -> None:
        nonlocal y
        font = "helv-bold" if bold else "helv"
        page.insert_text(fitz.Point(42, y), text, fontsize=size, fontname=font)
        y += size * 1.6

    _line("Governing Source Report", 14, bold=True)
    _line(f"Project: {data.get('project_id', '-')}", 9)
    _line(f"Selection method: {data.get('selection_method', '-')}", 9)
    _line(f"Rationale: {data.get('rationale', '-')}", 9)
    y += 8

    gf = data.get("governing_file") or {}
    if gf:
        _line("Governing File", 11, bold=True)
        for key, val in gf.items():
            _line(f"  {key}: {val}", 9)
    y += 6

    candidates = data.get("candidates", [])
    if candidates:
        _line("All Candidates (ranked)", 11, bold=True)
        for c in candidates:
            _line(f"  • {c.get('original_filename', '-')}  [{c.get('file_type', '-')}]"
                  f"  confidence={c.get('classification_confidence', '-')}", 9)

    doc.save(path)
    doc.close()


def _write_completeness_matrix_pdf(path: Path, data: dict) -> None:
    import fitz
    doc = fitz.open()
    page = doc.new_page(width=595, height=842)
    y = 42.0

    def _line(text: str, size: float = 9, bold: bool = False, color=(0, 0, 0)) -> None:
        nonlocal y
        font = "helv-bold" if bold else "helv"
        page.insert_text(fitz.Point(42, y), text, fontsize=size, fontname=font, color=color)
        y += size * 1.6

    gate = data.get("gate_status", "?")
    gate_color = (0, 0.5, 0) if gate == "PASS" else (0.8, 0, 0)
    _line("Completeness Matrix Report", 14, bold=True)
    _line(f"Gate: {gate}  |  Overall: {data.get('overall', '-')}", 11, bold=True, color=gate_color)
    y += 6

    for label, key in (("AB Readiness", "ab_readiness"), ("GA Readiness", "ga_readiness")):
        rd = data.get(key, {})
        status = rd.get("status", "-")
        clr = (0, 0.5, 0) if status == "READY" else (0.8, 0, 0)
        _line(f"{label} — {status}", 11, bold=True, color=clr)
        _line(f"  Required: {rd.get('required_count', 0)}   Present: {rd.get('present_count', 0)}"
              f"   Missing: {rd.get('missing_count', 0)}", 9)
        if rd.get("missing_codes"):
            _line(f"  Missing: {', '.join(rd['missing_codes'])}", 8, color=(0.7, 0, 0))
        y += 4

    blockers = data.get("blocking_issues", [])
    if blockers:
        _line("Blocking Issues", 11, bold=True, color=(0.8, 0, 0))
        for b in blockers:
            affects = "/".join(b.get("affects", []))
            _line(f"  • {b['field_code']}  [{affects}]  — {b.get('note', '')}", 8)

    warnings = data.get("non_blocking_warnings", [])
    if warnings:
        y += 4
        _line("Non-Blocking Warnings", 10, bold=True, color=(0.6, 0.4, 0))
        for w in warnings:
            _line(f"  • {w.get('field_code', '-')}  — {w.get('note', '')}", 8)

    hr = data.get("human_review_required", [])
    if hr:
        y += 4
        _line("Human Review Required", 10, bold=True)
        _line("  " + ", ".join(hr), 8)

    doc.save(path)
    doc.close()


# --- DOCX builders ---

def _write_file_inventory_docx(path: Path, data: dict) -> None:
    from docx import Document
    doc = Document()
    doc.add_heading("File Inventory Report", 0)
    doc.add_paragraph(f"Project: {data.get('project_id', '-')}")
    doc.add_paragraph(
        f"Total: {data.get('total_files', 0)}  Parsed: {data.get('parsed_files', 0)}"
        f"  Failed: {data.get('failed_files', 0)}"
    )
    doc.add_heading("Files", level=1)
    table = doc.add_table(rows=1, cols=6)
    headers = ["Filename", "Type", "Category", "Source", "Role", "Status"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for f in data.get("files", []):
        cells = table.add_row().cells
        cells[0].text = f.get("original_filename", "-")
        cells[1].text = f.get("file_type", "-")
        cells[2].text = f.get("file_category", "-")
        cells[3].text = f.get("source_application", "-")
        cells[4].text = f.get("likely_role", "-")
        cells[5].text = f.get("processing_status", "-")
    doc.save(path)


def _write_governing_source_docx(path: Path, data: dict) -> None:
    from docx import Document
    doc = Document()
    doc.add_heading("Governing Source Report", 0)
    doc.add_paragraph(f"Project: {data.get('project_id', '-')}")
    doc.add_paragraph(f"Method: {data.get('selection_method', '-')}")
    doc.add_paragraph(f"Rationale: {data.get('rationale', '-')}")
    gf = data.get("governing_file") or {}
    if gf:
        doc.add_heading("Governing File", level=1)
        for key, val in gf.items():
            doc.add_paragraph(f"{key}: {val}")
    candidates = data.get("candidates", [])
    if candidates:
        doc.add_heading("All Candidates", level=1)
        table = doc.add_table(rows=1, cols=3)
        for i, h in enumerate(["Filename", "Type", "Confidence"]):
            table.rows[0].cells[i].text = h
        for c in candidates:
            cells = table.add_row().cells
            cells[0].text = c.get("original_filename", "-")
            cells[1].text = c.get("file_type", "-")
            cells[2].text = str(c.get("classification_confidence", "-"))
    doc.save(path)


def _write_completeness_matrix_docx(path: Path, data: dict) -> None:
    from docx import Document
    doc = Document()
    doc.add_heading("Completeness Matrix Report", 0)
    doc.add_paragraph(f"Gate: {data.get('gate_status', '-')}  |  Overall: {data.get('overall', '-')}")
    for label, key in (("AB Readiness", "ab_readiness"), ("GA Readiness", "ga_readiness")):
        rd = data.get(key, {})
        doc.add_heading(label, level=1)
        doc.add_paragraph(
            f"Status: {rd.get('status', '-')}  Required: {rd.get('required_count', 0)}"
            f"  Present: {rd.get('present_count', 0)}  Missing: {rd.get('missing_count', 0)}"
        )
        if rd.get("missing_codes"):
            doc.add_paragraph("Missing: " + ", ".join(rd["missing_codes"]))
    blockers = data.get("blocking_issues", [])
    if blockers:
        doc.add_heading("Blocking Issues", level=1)
        table = doc.add_table(rows=1, cols=3)
        for i, h in enumerate(["Field Code", "Affects", "Note"]):
            table.rows[0].cells[i].text = h
        for b in blockers:
            cells = table.add_row().cells
            cells[0].text = b.get("field_code", "-")
            cells[1].text = "/".join(b.get("affects", []))
            cells[2].text = b.get("note", "")
    doc.save(path)


_PDF_BUILDERS = {
    "file_inventory.json": _write_file_inventory_pdf,
    "governing_source.json": _write_governing_source_pdf,
    "completeness_matrix.json": _write_completeness_matrix_pdf,
}

_DOCX_BUILDERS = {
    "file_inventory.json": _write_file_inventory_docx,
    "governing_source.json": _write_governing_source_docx,
    "completeness_matrix.json": _write_completeness_matrix_docx,
}


def _pdf_to_png(pdf_path: Path, png_path: Path) -> None:
    import fitz
    doc = fitz.open(str(pdf_path))
    page = doc[0]
    pix = page.get_pixmap(dpi=150)
    pix.save(str(png_path))
    doc.close()


# ---------------------------------------------------------------------------
# Processed/ endpoints
# ---------------------------------------------------------------------------

@router.get("/{project_id}/processed/files")
async def list_processed_files(project_id: str):
    try:
        pid = UUID(project_id)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail="Invalid project_id") from exc

    processed_root = _project_processed_root(pid)
    files = []
    if processed_root.exists():
        files = [
            {
                "name": path.name,
                "size_bytes": path.stat().st_size,
                "modified_at": datetime.fromtimestamp(path.stat().st_mtime).isoformat(),
            }
            for path in sorted(processed_root.iterdir())
            if path.is_file() and path.suffix == ".json"
        ]
    return {"project_id": project_id, "files": files}


@router.get("/{project_id}/processed/download/{filename}")
async def download_processed_file(project_id: str, filename: str):
    try:
        pid = UUID(project_id)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail="Invalid project_id") from exc

    processed_root = _project_processed_root(pid).resolve()
    target = (processed_root / filename).resolve()
    try:
        target.relative_to(processed_root)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail="Invalid filename") from exc
    if not target.exists() or not target.is_file():
        raise HTTPException(status_code=404, detail=f"File '{filename}' not found")
    return FileResponse(target, filename=target.name, media_type="application/json")


@router.get("/{project_id}/processed/export/{filename}/{format_name}")
async def export_processed_file(project_id: str, filename: str, format_name: str):
    try:
        pid = UUID(project_id)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail="Invalid project_id") from exc

    fmt = format_name.lower()
    if fmt not in {"json", "pdf", "docx", "png"}:
        raise HTTPException(status_code=400, detail="Supported formats: json, pdf, docx, png")

    data = _load_processed_json(pid, filename)
    stem = Path(filename).stem
    export_dir = _processed_export_dir(pid)

    if fmt == "json":
        out = export_dir / f"{stem}.json"
        out.write_text(json.dumps(data, indent=2, default=str), encoding="utf-8")
        return FileResponse(out, filename=out.name, media_type="application/json")

    if fmt == "pdf":
        builder = _PDF_BUILDERS.get(filename)
        if builder is None:
            out = export_dir / f"{stem}.pdf"
            import fitz
            doc = fitz.open()
            page = doc.new_page(width=595, height=842)
            text = json.dumps(data, indent=2, default=str)
            page.insert_textbox(fitz.Rect(42, 42, 553, 800), text,
                                fontsize=8, fontname="helv", lineheight=1.3)
            doc.save(str(out))
            doc.close()
        else:
            out = export_dir / f"{stem}.pdf"
            builder(out, data)
        return FileResponse(out, filename=out.name, media_type="application/pdf")

    if fmt == "docx":
        builder = _DOCX_BUILDERS.get(filename)
        if builder is None:
            raise HTTPException(
                status_code=400,
                detail=f"DOCX export not supported for '{filename}'. "
                       "Supported files: file_inventory.json, governing_source.json, "
                       "completeness_matrix.json"
            )
        out = export_dir / f"{stem}.docx"
        builder(out, data)
        media = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        return FileResponse(out, filename=out.name, media_type=media)

    # PNG — render PDF page 1 to image
    pdf_builder = _PDF_BUILDERS.get(filename)
    pdf_out = export_dir / f"{stem}_tmp.pdf"
    if pdf_builder is None:
        import fitz as _fitz
        doc = _fitz.open()
        page = doc.new_page(width=595, height=842)
        text = json.dumps(data, indent=2, default=str)
        page.insert_textbox(_fitz.Rect(42, 42, 553, 800), text,
                            fontsize=8, fontname="helv", lineheight=1.3)
        doc.save(str(pdf_out))
        doc.close()
    else:
        pdf_builder(pdf_out, data)
    png_out = export_dir / f"{stem}.png"
    _pdf_to_png(pdf_out, png_out)
    return FileResponse(png_out, filename=png_out.name, media_type="image/png")


# ---------------------------------------------------------------------------
# Benchmark Evidence Exports
# ---------------------------------------------------------------------------

@router.get("/{project_id}/export/audit")
async def export_audit_log(project_id: str, db: Session = Depends(get_db)):
    """Export the full immutable audit trail for a project."""
    from app.db.models import AuditEventLog
    try:
        pid = UUID(project_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid project_id")
    
    logs = db.query(AuditEventLog).filter(AuditEventLog.project_id == pid).order_by(AuditEventLog.created_at).all()
    data = [
        {
            "id": str(l.id),
            "event_type": l.event_type,
            "actor": l.actor,
            "stage_code": l.stage_code,
            "field_code": l.field_code,
            "detail": json.loads(l.detail or "{}"),
            "created_at": l.created_at.isoformat()
        }
        for l in logs
    ]
    
    export_dir = _processed_export_dir(pid)
    out = export_dir / "audit_log.json"
    out.write_text(json.dumps(data, indent=2, default=str), encoding="utf-8")
    return FileResponse(out, filename=out.name, media_type="application/json")


@router.get("/{project_id}/audit-events")
async def get_audit_events(
    project_id: str,
    limit: int = 200,
    db: Session = Depends(get_db),
):
    """Return runtime audit events for operator visibility in the UI."""
    try:
        pid = UUID(project_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid project_id")

    safe_limit = max(1, min(limit, 1000))
    rows = (
        db.query(AuditEventLog)
        .filter(AuditEventLog.project_id == pid)
        .order_by(AuditEventLog.created_at.desc())
        .limit(safe_limit)
        .all()
    )
    events = [
        {
            "id": str(r.id),
            "event_type": r.event_type,
            "actor": r.actor,
            "stage_code": r.stage_code,
            "field_code": r.field_code,
            "detail": json.loads(r.detail or "{}"),
            "created_at": r.created_at.isoformat() if r.created_at else None,
        }
        for r in rows
    ]
    return {"project_id": str(pid), "count": len(events), "events": events}


@router.get("/{project_id}/export/validation")
async def export_validation_results(project_id: str, db: Session = Depends(get_db)):
    """Export all per-field validation records (PRESENT/MISSING/SUSPICIOUS)."""
    from app.db.models import ValidationResult
    try:
        pid = UUID(project_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid project_id")
    
    items = db.query(ValidationResult).filter(ValidationResult.project_id == pid).order_by(ValidationResult.created_at).all()
    data = [
        {
            "id": str(i.id),
            "stage_code": i.stage_code,
            "field_code": i.field_code,
            "status": i.status,
            "severity": i.severity,
            "source": i.source,
            "value": i.value,
            "note": i.note,
            "created_at": i.created_at.isoformat()
        }
        for i in items
    ]
    
    export_dir = _processed_export_dir(pid)
    out = export_dir / "validation_results.json"
    out.write_text(json.dumps(data, indent=2, default=str), encoding="utf-8")
    return FileResponse(out, filename=out.name, media_type="application/json")


@router.get("/{project_id}/export/blockers")
async def export_blockers(project_id: str, db: Session = Depends(get_db)):
    """Export a consolidated list of critical blockers (MISSING mandatory fields)."""
    from app.db.models import ValidationResult
    try:
        pid = UUID(project_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid project_id")
    
    # Blockers are MISSING status OR CRITICAL severity
    items = db.query(ValidationResult).filter(
        ValidationResult.project_id == pid,
        (ValidationResult.status == "MISSING") | (ValidationResult.severity == "CRITICAL")
    ).all()
    
    data = [
        {
            "stage_code": i.stage_code,
            "field_code": i.field_code,
            "status": i.status,
            "severity": i.severity,
            "note": i.note
        }
        for i in items
    ]
    
    export_dir = _processed_export_dir(pid)
    out = export_dir / "blockers.json"
    out.write_text(json.dumps(data, indent=2, default=str), encoding="utf-8")
    return FileResponse(out, filename=out.name, media_type="application/json")


@router.get("/{project_id}/export/status")
async def export_stage_status(project_id: str, db: Session = Depends(get_db)):
    """Export the full stage status history and current state."""
    from app.orchestrator.controller import OrchestrationController
    try:
        pid = UUID(project_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid project_id")
    
    data = OrchestrationController().get_pipeline_status(project_id, db=db)
    
    export_dir = _processed_export_dir(pid)
    out = export_dir / "stage_status.json"
    out.write_text(json.dumps(data, indent=2, default=str), encoding="utf-8")
    return FileResponse(out, filename=out.name, media_type="application/json")


@router.get("/{project_id}/export/excel")
async def export_evidence_excel(project_id: str, db: Session = Depends(get_db)):
    """Export all project evidence into a single multi-sheet Excel file."""
    import pandas as pd
    from io import BytesIO
    
    try:
        pid = UUID(project_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid project_id")
    
    project = db.query(Project).filter(Project.id == pid).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    # Helper to convert query to DataFrame
    def q_to_df(query):
        return pd.DataFrame(_query_rows(query))

    # 1. Project Summary
    df_proj = pd.DataFrame([{
        "Project ID": str(project.id),
        "Proposal ID": project.proposal_id,
        "Name": project.name,
        "Status": project.status,
        "Created At": project.created_at,
    }])

    # 2. File Inventory
    df_files = q_to_df(db.query(ProjectFile).filter(ProjectFile.project_id == pid))
    
    # 3. Extraction Log (Parser Runs)
    df_parsers = q_to_df(db.query(ParserRun).filter(ParserRun.project_id == pid))
    
    # 4. Field Values
    df_fields = q_to_df(db.query(ExtractedFieldValue).filter(ExtractedFieldValue.project_id == pid))
    
    # 5. Validation Results
    df_val = q_to_df(db.query(ValidationResult).filter(ValidationResult.project_id == pid))
    
    # 6. Blockers
    df_blockers = df_val[df_val["status"].isin(["MISSING", "SUSPICIOUS"]) | (df_val["severity"] == "CRITICAL")]
    
    # 7. Stage Status
    df_stages = q_to_df(db.query(Stage).filter(Stage.project_id == pid))
    
    # 8. Audit Log
    df_audit = q_to_df(db.query(AuditEventLog).filter(AuditEventLog.project_id == pid).order_by(AuditEventLog.created_at))

    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        df_proj.to_excel(writer, sheet_name="intake", index=False)
        df_files.to_excel(writer, sheet_name="file_registry", index=False)
        df_parsers.to_excel(writer, sheet_name="extraction_log", index=False)
        df_fields.to_excel(writer, sheet_name="field_value_store", index=False)
        df_val.to_excel(writer, sheet_name="validation_results", index=False)
        df_blockers.to_excel(writer, sheet_name="blockers", index=False)
        df_stages.to_excel(writer, sheet_name="stage_status", index=False)
        df_audit.to_excel(writer, sheet_name="audit_sample", index=False)
    
    output.seek(0)
    filename = f"{project.name.replace(' ', '_')}_Evidence_Export.xlsx"
    export_dir = _processed_export_dir(pid)
    out_file = export_dir / filename
    out_file.write_bytes(output.getvalue())
    
    return FileResponse(
        out_file, 
        filename=filename, 
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )


@router.get("/{project_id}/export/evidence-files")
async def export_evidence_files(project_id: str, db: Session = Depends(get_db)):
    """
    Export benchmark-aligned evidence JSON files into processed_exports/.
    Writes: intake, file_registry, extraction_log, field_value_store,
            validation_results, blockers, stage_status, audit_sample,
            defect_log, integrity_manifest.
    """
    import hashlib

    try:
        pid = UUID(project_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid project_id")

    project = db.query(Project).filter(Project.id == pid).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    intake = [{
        "project_id": str(project.id),
        "proposal_id": project.proposal_id,
        "name": project.name,
        "location": project.location,
        "project_type": project.project_type,
        "status": str(project.status),
        "created_at": project.created_at,
    }]
    file_registry     = _query_rows(db.query(ProjectFile).filter(ProjectFile.project_id == pid))
    extraction_log    = _query_rows(db.query(ParserRun).filter(ParserRun.project_id == pid))
    field_value_store = _query_rows(db.query(ExtractedFieldValue).filter(ExtractedFieldValue.project_id == pid))
    validation_results = _query_rows(db.query(ValidationResult).filter(ValidationResult.project_id == pid))
    blockers = [
        row for row in validation_results
        if row.get("status") in {"MISSING", "SUSPICIOUS"} or row.get("severity") == "CRITICAL"
    ]
    stage_status  = _query_rows(db.query(Stage).filter(Stage.project_id == pid))
    audit_sample  = _query_rows(
        db.query(AuditEventLog).filter(AuditEventLog.project_id == pid).order_by(AuditEventLog.created_at)
    )

    # Defect log — escalations + correction events
    raw_escalations = _query_rows(db.query(Escalation).filter(Escalation.project_id == pid))
    raw_corrections = _query_rows(db.query(CorrectionEvent).filter(CorrectionEvent.project_id == pid))
    defect_log = {
        "summary": {
            "total_escalations": len(raw_escalations),
            "open":     sum(1 for e in raw_escalations if e.get("status") == "OPEN"),
            "resolved": sum(1 for e in raw_escalations if e.get("status") == "RESOLVED"),
            "critical": sum(1 for e in raw_escalations if e.get("severity") == "CRITICAL"),
            "total_corrections": len(raw_corrections),
        },
        "escalations": raw_escalations,
        "corrections": raw_corrections,
    }

    export_dir = _processed_export_dir(pid)
    payload_map = {
        "intake.json":            intake,
        "file_registry.json":     file_registry,
        "extraction_log.json":    extraction_log,
        "field_value_store.json": field_value_store,
        "validation_results.json": validation_results,
        "blockers.json":          blockers,
        "stage_status.json":      stage_status,
        "audit_sample.json":      audit_sample,
        "defect_log.json":        defect_log,
    }
    for filename, payload in payload_map.items():
        _write_json(export_dir / filename, payload)

    # Integrity manifest — SHA-256 of every exported file
    manifest: dict[str, dict] = {}
    for filename in payload_map:
        f = export_dir / filename
        if f.exists():
            digest = hashlib.sha256(f.read_bytes()).hexdigest()
            manifest[filename] = {"sha256": digest, "size_bytes": f.stat().st_size}
    _write_json(export_dir / "integrity_manifest.json", manifest)

    all_files = sorted(payload_map.keys()) + ["integrity_manifest.json"]
    return {
        "project_id": str(pid),
        "export_dir": str(export_dir),
        "files": all_files,
    }


@router.get("/{project_id}/export/pipeline-snapshot-png")
async def export_pipeline_snapshot_png(project_id: str, db: Session = Depends(get_db)):
    """Render a PNG snapshot of the current pipeline status for BM-001 evidence capture."""
    import fitz
    from app.orchestrator.controller import OrchestrationController

    try:
        pid = UUID(project_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid project_id")

    project = db.query(Project).filter(Project.id == pid).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    status_data = OrchestrationController().get_pipeline_status(project_id, db=db)

    _STATUS_COLOR = {
        "PASSED":             (0.0, 0.55, 0.27),
        "PASS_WITH_WARNINGS": (0.55, 0.45, 0.0),
        "BLOCKED":            (0.80, 0.30, 0.0),
        "FAILED":             (0.75, 0.0,  0.0),
        "RUNNING":            (0.0,  0.40, 0.80),
        "PENDING":            (0.50, 0.50, 0.50),
    }

    doc  = fitz.open()
    page = doc.new_page(width=800, height=700)
    y    = 36.0

    def _text(txt: str, size: float = 10, bold: bool = False,
               color: tuple = (0, 0, 0), x: float = 36) -> None:
        nonlocal y
        page.insert_text(
            fitz.Point(x, y), txt,
            fontsize=size, fontname="hebo" if bold else "helv", color=color,
        )
        y += size * 1.8

    # ── Header ────────────────────────────────────────────────────────────────
    _text("BM-001 Pipeline Status Snapshot", 16, bold=True)
    _text(f"Project: {project.name}  |  {project.proposal_id}", 10)
    _text(
        f"Captured: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S UTC')}",
        9, color=(0.4, 0.4, 0.4),
    )
    y += 12

    # ── Stage cards ───────────────────────────────────────────────────────────
    _text("Pipeline Stages", 12, bold=True)
    y += 6

    for stage in status_data.get("stages", []):
        sc     = stage["stage_code"]
        st     = stage["status"]
        color  = _STATUS_COLOR.get(st, (0.5, 0.5, 0.5))
        card_y = y

        # Card background
        page.draw_rect(fitz.Rect(36, card_y, 764, card_y + 30),
                       color=(0.93, 0.93, 0.93), fill=(0.93, 0.93, 0.93))
        # Status badge
        page.draw_rect(fitz.Rect(36, card_y, 106, card_y + 30),
                       color=color, fill=color)
        page.insert_text(fitz.Point(41, card_y + 19), sc,
                         fontsize=9, fontname="hebo", color=(1, 1, 1))
        # Status label
        page.insert_text(fitz.Point(114, card_y + 19), st,
                         fontsize=9, fontname="hebo", color=color)
        # Detail (error or overall result)
        err    = stage.get("error_message") or ""
        result = stage.get("result") or {}
        detail = err or result.get("overall") or result.get("status") or ""
        if detail:
            page.insert_text(fitz.Point(260, card_y + 19), str(detail)[:80],
                             fontsize=8, fontname="helv", color=(0.3, 0.3, 0.3))
        y += 34

    y += 14

    # ── Gate checkpoints ──────────────────────────────────────────────────────
    checkpoints = status_data.get("checkpoints", [])
    if checkpoints:
        _text("Gate Checkpoints", 12, bold=True)
        y += 4
        for cp in checkpoints:
            gs      = cp.get("gate_status", "?")
            gc      = (0.0, 0.55, 0.27) if gs == "PASS" else (0.75, 0.0, 0.0)
            label   = f"[{cp.get('stage_code', '')}]  {cp.get('checkpoint_label', '')}  —  {gs}"
            _text(label[:100], 8, color=gc)

    # ── Footer ────────────────────────────────────────────────────────────────
    page.insert_text(
        fitz.Point(36, 688),
        "Steel Detailing Automation — BM-001 Evidence Pack  |  v2.0",
        fontsize=7, fontname="helv", color=(0.6, 0.6, 0.6),
    )

    export_dir = _processed_export_dir(pid)
    pdf_tmp    = export_dir / "pipeline_snapshot_tmp.pdf"
    png_out    = export_dir / "pipeline_snapshot.png"
    doc.save(str(pdf_tmp))
    doc.close()
    _pdf_to_png(pdf_tmp, png_out)

    return FileResponse(png_out, filename=png_out.name, media_type="image/png")
