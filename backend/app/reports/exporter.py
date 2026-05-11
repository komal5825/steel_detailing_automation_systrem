"""Export report payloads to JSON, XLSX, and DOCX formats."""
from __future__ import annotations

import json
from pathlib import Path


# ---------------------------------------------------------------------------
# JSON
# ---------------------------------------------------------------------------

def export_json(report: dict, out_path: Path) -> Path:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(json.dumps(report, indent=2, default=str), encoding="utf-8")
    return out_path


# ---------------------------------------------------------------------------
# XLSX (multi-sheet)
# ---------------------------------------------------------------------------

def export_xlsx(report: dict, out_path: Path) -> Path:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment

    out_path.parent.mkdir(parents=True, exist_ok=True)
    wb = Workbook()
    wb.remove(wb.active)  # remove default sheet

    report_type = report.get("header", {}).get("report_type", "REPORT")

    def _header_row(ws, cols: list[str]) -> None:
        ws.append(cols)
        for cell in ws[1]:
            cell.font      = Font(bold=True, color="FFFFFF")
            cell.fill      = PatternFill("solid", fgColor="1E3A5F")
            cell.alignment = Alignment(horizontal="center")

    def _add_kv_sheet(name: str, data: dict) -> None:
        ws = wb.create_sheet(name[:31])
        _header_row(ws, ["Field", "Value"])
        for k, v in data.items():
            ws.append([str(k), str(v) if not isinstance(v, (list, dict)) else json.dumps(v, default=str)])
        ws.column_dimensions["A"].width = 28
        ws.column_dimensions["B"].width = 45

    # ── Common: Header ───────────────────────────────────────────────────────
    _add_kv_sheet("Header", report.get("header", {}))

    # ── Daily-specific sheets ─────────────────────────────────────────────────
    if report_type == "DAILY":
        _add_kv_sheet("Execution Summary", report.get("execution_summary", {}))

        # Lane Summary
        ws_lane = wb.create_sheet("Lane Summary")
        _header_row(ws_lane, ["Lane", "Stages", "Statuses", "Overall"])
        for lane, info in report.get("lane_summary", {}).items():
            ws_lane.append([lane, ", ".join(info["stages"]), ", ".join(info["statuses"]), info["overall"]])

        # Gate Snapshot
        ws_gate = wb.create_sheet("Gate Snapshot")
        _header_row(ws_gate, ["Gate", "Stage", "Checkpoint", "Status", "Recorded At"])
        for label, info in report.get("gate_snapshot", {}).items():
            ws_gate.append([label, info["stage_code"], info["checkpoint_label"], info["gate_status"], info.get("created_at", "")])

        # Severity Summary
        _add_kv_sheet("Severity Summary", report.get("severity_summary", {}))

        # Top Blockers
        ws_bl = wb.create_sheet("Top Blockers")
        _header_row(ws_bl, ["Field Code", "Stage", "Status", "Severity", "Note", "Source"])
        for b in report.get("top_blockers", []):
            ws_bl.append([b.get("field_code"), b.get("stage_code"), b.get("status"), b.get("severity"), b.get("note"), b.get("source")])

        # Manual Review
        ws_mr = wb.create_sheet("Manual Review")
        mr = report.get("manual_review_summary", {})
        ws_mr.append(["Total", "Open", "Resolved"])
        ws_mr.append([mr.get("total"), mr.get("open"), mr.get("resolved")])
        if mr.get("open_items"):
            ws_mr.append([])
            ws_mr.append(["ID", "Stage", "Severity", "Reason"])
            for item in mr["open_items"]:
                ws_mr.append([item.get("id"), item.get("stage_code"), item.get("severity"), item.get("reason")])

        # Stage Detail
        ws_st = wb.create_sheet("Stage Detail")
        _header_row(ws_st, ["Stage", "Status", "Revision", "Started At", "Completed At", "Error"])
        for s in report.get("stage_detail", []):
            ws_st.append([s["stage_code"], s["status"], s["revision"], s["started_at"], s["completed_at"], s.get("error_message")])

        # Defect Summary
        _add_kv_sheet("Defect Summary", report.get("defect_summary", {}))

        # Next-Day Critical Path
        ncp = report.get("next_day_critical_path", {})
        ws_ncp = wb.create_sheet("Next-Day Critical")
        ws_ncp.append(["Action Required", str(ncp.get("action_required", False))])
        ws_ncp.append([])
        ws_ncp.append(["Blocked Stages"])
        ws_ncp.append(["Stage Code", "Status", "Error"])
        for s in ncp.get("blocked_stages", []):
            ws_ncp.append([s.get("stage_code"), s.get("status"), s.get("error")])
        ws_ncp.append([])
        ws_ncp.append(["Critical Missing Fields"])
        ws_ncp.append(["Field Code", "Stage", "Note"])
        for f in ncp.get("critical_missing", []):
            ws_ncp.append([f.get("field_code"), f.get("stage_code"), f.get("note")])

    # ── Weekly-specific sheets ────────────────────────────────────────────────
    elif report_type == "WEEKLY":
        _add_kv_sheet("Throughput", report.get("throughput", {}))
        _add_kv_sheet("Gate Trend", report.get("gate_trend", {}))
        _add_kv_sheet("Validation Summary", report.get("validation_summary", {}))

        ws_rb = wb.create_sheet("Recurring Blockers")
        _header_row(ws_rb, ["Field Code", "Occurrences"])
        for rb in report.get("recurring_blockers", []):
            ws_rb.append([rb["field_code"], rb["occurrences"]])

        _add_kv_sheet("Parser Stability", report.get("parser_stability", {}))
        _add_kv_sheet("Defect Summary", report.get("defect_summary", {}))

        ws_imp = wb.create_sheet("Improvement Areas")
        _header_row(ws_imp, ["#", "Improvement Area"])
        for i, item in enumerate(report.get("top_10_improvement_areas", []), 1):
            ws_imp.append([i, item])

    # ── Monthly-specific sheets ───────────────────────────────────────────────
    elif report_type == "MONTHLY":
        _add_kv_sheet("System Effectiveness", report.get("system_effectiveness", {}))
        _add_kv_sheet("Component Health", report.get("component_health", {}))
        _add_kv_sheet("Pipeline Overview", report.get("pipeline_overview", {}))

        rca = report.get("root_cause_analysis", {})
        ws_rca = wb.create_sheet("Root Cause Analysis")
        _header_row(ws_rca, ["Field Code", "Failure Count"])
        for item in rca.get("top_failing_fields", []):
            ws_rca.append([item["field_code"], item["count"]])
        ws_rca.append([])
        ws_rca.append(["Source", "Failure Count"])
        for item in rca.get("top_failing_sources", []):
            ws_rca.append([item["source"], item["count"]])

        _add_kv_sheet("Benchmark Evidence", report.get("benchmark_and_evidence", {}))
        _add_kv_sheet("Defect Summary", report.get("defect_summary", {}))

        ws_ir = wb.create_sheet("Improvement Register")
        _header_row(ws_ir, ["Area", "Item", "Priority"])
        for item in report.get("improvement_register", []):
            ws_ir.append([item.get("area"), item.get("item"), item.get("priority")])

        _add_kv_sheet("Management Summary", report.get("management_summary", {}))

    wb.save(out_path)
    return out_path


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def export_docx(report: dict, out_path: Path) -> Path:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    header    = report.get("header", {})
    rtype     = header.get("report_type", "REPORT")
    proj_name = header.get("project_name", "")
    gen_at    = header.get("generated_at", "")

    # Title
    title = doc.add_heading(f"Infiniti Steel Detailing — {rtype} Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Project: {proj_name}  |  Generated: {gen_at}")
    doc.add_paragraph(f"Proposal ID: {header.get('proposal_id', '—')}  |  Build: {header.get('build_version', '2.0')}")
    doc.add_paragraph("")

    def _section(title: str, data: dict) -> None:
        doc.add_heading(title, level=1)
        if not data:
            doc.add_paragraph("No data.")
            return
        tbl = doc.add_table(rows=1, cols=2)
        tbl.rows[0].cells[0].text = "Field"
        tbl.rows[0].cells[1].text = "Value"
        for k, v in data.items():
            row = tbl.add_row().cells
            row[0].text = str(k)
            row[1].text = str(v) if not isinstance(v, (list, dict)) else json.dumps(v, default=str)

    def _list_section(title: str, items: list, cols: list[str], row_fn) -> None:
        doc.add_heading(title, level=1)
        if not items:
            doc.add_paragraph("None.")
            return
        tbl = doc.add_table(rows=1, cols=len(cols))
        for i, col in enumerate(cols):
            tbl.rows[0].cells[i].text = col
        for item in items:
            cells = tbl.add_row().cells
            values = row_fn(item)
            for i, val in enumerate(values):
                cells[i].text = str(val) if val is not None else ""

    if rtype == "DAILY":
        _section("Execution Summary", report.get("execution_summary", {}))
        _section("Severity Summary", report.get("severity_summary", {}))
        _list_section(
            "Top Blockers",
            report.get("top_blockers", []),
            ["Field Code", "Stage", "Status", "Severity", "Note"],
            lambda b: [b.get("field_code"), b.get("stage_code"), b.get("status"), b.get("severity"), b.get("note")],
        )
        _section("Defect Summary", report.get("defect_summary", {}))
        _section("Source / Fallback Summary", report.get("source_fallback_summary", {}))
        ncp = report.get("next_day_critical_path", {})
        doc.add_heading("Next-Day Critical Path", level=1)
        doc.add_paragraph(f"Action Required: {ncp.get('action_required', False)}")
        if ncp.get("blocked_stages"):
            doc.add_paragraph("Blocked Stages: " + ", ".join(s["stage_code"] for s in ncp["blocked_stages"]))
        if ncp.get("critical_missing"):
            doc.add_paragraph("Critical Missing Fields: " + ", ".join(f["field_code"] for f in ncp["critical_missing"]))

    elif rtype == "WEEKLY":
        _section("Throughput", report.get("throughput", {}))
        _section("Gate Trend", report.get("gate_trend", {}))
        _section("Validation Summary", report.get("validation_summary", {}))
        _section("Parser Stability", report.get("parser_stability", {}))
        _section("Defect Summary", report.get("defect_summary", {}))
        doc.add_heading("Top Improvement Areas", level=1)
        for i, item in enumerate(report.get("top_10_improvement_areas", []), 1):
            doc.add_paragraph(f"{i}. {item}", style="List Number")

    elif rtype == "MONTHLY":
        _section("System Effectiveness", report.get("system_effectiveness", {}))
        _section("Component Health", report.get("component_health", {}))
        _section("Pipeline Overview", report.get("pipeline_overview", {}))
        _section("Benchmark & Evidence", report.get("benchmark_and_evidence", {}))
        _section("Defect Summary", report.get("defect_summary", {}))
        doc.add_heading("Improvement Register", level=1)
        for item in report.get("improvement_register", []):
            doc.add_paragraph(f"[{item.get('priority')}] {item.get('area')}: {item.get('item')}", style="List Bullet")
        mgmt = report.get("management_summary", {})
        doc.add_heading("Management Summary", level=1)
        doc.add_paragraph(f"Release Ready: {mgmt.get('release_ready')}  |  Effectiveness Score: {mgmt.get('effectiveness_score')}")
        if mgmt.get("key_risks"):
            doc.add_paragraph("Key Risks:")
            for risk in mgmt["key_risks"]:
                doc.add_paragraph(risk, style="List Bullet")

    doc.save(out_path)
    return out_path
